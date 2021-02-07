import wikipedia
from docx import Document

wikipedia.set_lang('th')
product = ['ทุเรียน','ลำไย','แอปเปิ้ล','มังคุด']

document = Document()
document.add_heading('รายการผลไม้', 0)

for pd in product:
	#content = wikipedia.summary(pd) #สรุป
	content = wikipedia.page(pd).content
	document.add_heading(pd, level=1)
	p = document.add_paragraph('')
	p.add_run(content + '\n\n\n')
	#document.add_page_break()

document.save('fruit.docx')